from datetime import date
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "AI_Study_Assistant_Progress_and_Roadmap.docx"
NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "247A52"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_status_table(doc):
    rows = [
        ("Account and security", "Completed", "Registration, login, logout, secure server-side sessions, ownership checks, and profile editing."),
        ("Dashboard", "Completed", "Personal overview with topic activity, study statistics, and flashcard summary."),
        ("Topics", "Completed", "Create, view, search, edit, and delete study topics."),
        ("Notes", "Completed", "Create, search, paginate, edit, move between topics, and delete notes."),
        ("Document knowledge base", "Completed", "Upload PDF or text files, track processing status, list files, and delete files."),
        ("AI tutor", "Completed", "Topic-based chat, saved message history, clear history, source citations, and optional document-specific questions."),
        ("RAG retrieval", "Completed", "Hybrid vector and keyword retrieval over notes and uploaded documents, with fallback behavior."),
        ("Flashcards", "Completed", "Manual and AI-generated cards, editing, archiving, regeneration, deletion, and source links."),
        ("Spaced repetition", "Completed", "Due-card queue, review ratings, scheduling, deck statistics, and overall review summary."),
        ("Study history", "Completed", "Activity timeline, filters, pagination, and study statistics."),
        ("Production foundation", "Completed", "Database migrations, health/readiness checks, Docker setup, rate limiting, structured logs, and security headers."),
        ("Automated testing", "Completed", "Unit and integration coverage for core logic, authentication, CRUD, documents, AI chat, rate limits, and pagination."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].text = "Feature area"
    table.rows[0].cells[1].text = "Status"
    table.rows[0].cells[2].text = "What is available now"
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.bold = True
            run.font.size = Pt(9.5)
    for idx, item in enumerate(rows):
        cells = table.add_row().cells
        for col, value in enumerate(item):
            cells[col].text = value
            set_cell_margins(cells[col])
            cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx % 2:
                set_cell_shading(cells[col], LIGHT_GRAY)
            for p in cells[col].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(9)
        cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(GREEN)
        cells[1].paragraphs[0].runs[0].bold = True
    set_table_widths(table, [1.55, 0.85, 4.10])


def add_roadmap_table(doc):
    rows = [
        ("Phase 1: Release readiness", "Next", "Deploy the application; add CI/CD; use managed PostgreSQL with pgvector; configure backups, monitoring, and environment secrets.", "A stable version that can be demonstrated and tested by real users."),
        ("Phase 2: Storage and scale", "Next", "Replace local uploads with object storage and replace in-memory rate limits with a shared Redis or gateway-based limiter.", "Reliable files and consistent protection across multiple server instances."),
        ("Phase 3: Learning tools", "Planned", "Add quizzes generated from notes, configurable practice sessions, explanations for wrong answers, and progress by topic.", "More active recall and clearer evidence of learning progress."),
        ("Phase 4: Reminders and goals", "Planned", "Add study goals, review reminders, streaks based on real activity, and a weekly plan.", "Better study consistency and retention."),
        ("Phase 5: Reporting and export", "Planned", "Add downloadable notes, flashcards, study summaries, and supervisor/student progress reports.", "Portable learning materials and easier progress review."),
        ("Phase 6: Experience improvements", "Future", "Improve accessibility, responsive/mobile behavior, error recovery, and optional dark mode.", "A more inclusive and polished user experience."),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Roadmap stage", "Priority", "Planned addition", "Expected value"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], BLUE)
        set_cell_margins(table.rows[0].cells[i])
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col, value in enumerate(row):
            cells[col].text = value
            set_cell_margins(cells[col], top=90, bottom=90)
            cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx % 2:
                set_cell_shading(cells[col], LIGHT_BLUE)
            for p in cells[col].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8.5)
        cells[1].paragraphs[0].runs[0].bold = True
    set_table_widths(table, [1.15, 0.65, 2.85, 1.85])


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for name, size, before, after in (
    ("Heading 1", 16, 16, 8),
    ("Heading 2", 13, 12, 6),
    ("Heading 3", 12, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE if name != "Heading 3" else NAVY)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.text = "AI STUDY ASSISTANT  |  PROJECT PROGRESS REPORT"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in hp.runs:
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
add_page_number(section.footer.paragraphs[0])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("PROJECT PROGRESS & ROADMAP")
r.font.name = "Calibri"
r.font.size = Pt(23)
r.bold = True
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
r = p.add_run("AI Study Assistant (Studia)")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(MID_GRAY)

for label, value in (
    ("To", "Project Supervisor"),
    ("From", "AI Study Assistant Project Team"),
    ("Date", "27 July 2026"),
    ("Purpose", "Summary of completed features and proposed next development phases"),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(10)
p.paragraph_format.left_indent = Inches(0.18)
p.paragraph_format.right_indent = Inches(0.18)
p.paragraph_format.space_before = Pt(12)
r = p.add_run(
    "Current position: the project has progressed beyond a visual prototype into a working full-stack study platform. "
    "Its main learning workflows are implemented; the next priority is deployment hardening, user validation, and expansion of practice and reporting tools."
)
r.bold = True
r.font.color.rgb = RGBColor.from_string(NAVY)

doc.add_heading("1. Project overview", level=1)
doc.add_paragraph(
    "Studia is an AI-powered study assistant that helps students organize learning material, ask questions against their own notes and documents, create and review flashcards, and track study activity. "
    "The solution uses a modern web interface, a FastAPI backend, PostgreSQL, and retrieval-augmented generation (RAG) so AI answers can be grounded in the student's selected topic."
)

doc.add_heading("2. Features completed to date", level=1)
doc.add_paragraph(
    "The following capabilities are present in the current codebase. “Completed” means that both the user-facing flow and the supporting backend/API are implemented; final production deployment and user acceptance testing remain separate release activities."
)
add_status_table(doc)

doc.add_heading("3. Technical work completed", level=1)
add_bullet(doc, "Frontend: Responsive Next.js/React interface with dedicated pages for the dashboard, topics, notes, AI tutor, flashcards, study history, authentication, and settings.", "Frontend:")
add_bullet(doc, "Backend: Modular FastAPI service with clear separation between routes, services, repositories, database models, and validation schemas.", "Backend:")
add_bullet(doc, "Database: PostgreSQL schema managed through Alembic migrations, including users, sessions, topics, notes, AI messages, documents, vector chunks, sources, study history, and flashcards.", "Database:")
add_bullet(doc, "AI and retrieval: Support for Gemini, Groq, or OpenAI chat providers, with Gemini or OpenAI embeddings and 768-dimension pgvector storage.", "AI and retrieval:")
add_bullet(doc, "Reliability and security: Secure HTTP-only sessions, authentication and AI rate limits, request IDs, secret-redacting logs, CORS configuration, health checks, and ownership validation.", "Reliability and security:")
add_bullet(doc, "Quality assurance: Unit tests for algorithms and configuration plus integration tests for the principal API workflows.", "Quality assurance:")

doc.add_heading("4. Proposed features and improvements to add", level=1)
doc.add_paragraph(
    "The roadmap below is proposed from the current technical gaps and the learning goals of the product. The order can be adjusted after supervisor feedback and user testing."
)
add_roadmap_table(doc)

doc.add_heading("5. Immediate next milestone", level=1)
doc.add_paragraph(
    "The recommended next milestone is a supervised beta release. This milestone should focus on making the existing system dependable before adding many new features."
)
for text in (
    "Deploy one test environment with managed PostgreSQL and pgvector.",
    "Move uploaded files to persistent object storage.",
    "Add an automated build, test, migration, and deployment pipeline.",
    "Run end-to-end acceptance testing with a small group of students.",
    "Collect feedback on the AI tutor, flashcard review flow, and dashboard clarity.",
    "Use the feedback to select the first learning enhancement: quizzes or study goals/reminders.",
):
    add_bullet(doc, text)

doc.add_heading("6. Success criteria for the next review", level=1)
for text in (
    "A supervisor can register, create a topic, add notes, upload a document, and receive a source-grounded AI answer.",
    "A student can generate or create flashcards, complete a due-card review, and see updated statistics.",
    "All automated tests pass in the deployment pipeline.",
    "Uploaded documents remain available after application restarts.",
    "At least five student testers complete the main workflow and provide structured feedback.",
):
    add_bullet(doc, text)

doc.add_heading("7. Summary", level=1)
doc.add_paragraph(
    "The core product is now functionally broad: it combines content organization, document ingestion, grounded AI tutoring, spaced-repetition flashcards, and activity tracking in one system. "
    "The most valuable next step is to stabilize and validate this version in a real environment, then add quizzes, goals, reminders, and reporting based on observed student needs."
)

doc.core_properties.title = "AI Study Assistant - Project Progress and Roadmap"
doc.core_properties.subject = "Completed features and proposed next development phases"
doc.core_properties.author = "AI Study Assistant Project Team"
doc.core_properties.keywords = "AI study assistant, Studia, progress report, roadmap"
doc.save(OUT)
print(OUT)
