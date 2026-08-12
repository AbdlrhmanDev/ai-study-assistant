from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\Users\abdlr\Downloads\NTG_Traninig\Weekly_Report_Form_v2.docx")
OUTPUT = Path(r"C:\Users\abdlr\Downloads\NTG_Traninig\Project AI Study Assistant Website\Weekly_Report_AI_Study_Assistant_AR_EN.docx")


def set_bidi(paragraph, rtl=False):
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1" if rtl else "0")


def replace_cell(cell, arabic, english="", font_size=9):
    template_p = cell.paragraphs[0]
    template_rpr = None
    for run in template_p.runs:
        if run._r.rPr is not None:
            template_rpr = deepcopy(run._r.rPr)
            break
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2 if english else 0)
    p.paragraph_format.line_spacing = 1.0
    set_bidi(p, True)
    r = p.add_run(arabic)
    if template_rpr is not None:
        r._r.insert(0, deepcopy(template_rpr))
    r.font.size = Pt(font_size)
    r.font.bold = False
    if english:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        set_bidi(p2, False)
        r2 = p2.add_run(english)
        if template_rpr is not None:
            r2._r.insert(0, deepcopy(template_rpr))
        r2.font.size = Pt(font_size)
        r2.font.bold = False


def response_cell(row):
    # Narrative response rows are visually merged across all four grid columns.
    return row.cells[0]


doc = Document(SOURCE)
table = doc.tables[0]

# General information: retain unknown personal/organizational fields blank.
replace_cell(response_cell(table.rows[6]), "30 يوليو 2026", "30 July 2026", 9)
replace_cell(response_cell(table.rows[8]), "الأسبوع 31", "Week 31", 9)

# Project metadata. The form places Arabic/English labels at the outer edges.
metadata = {
    11: ("مساعد الدراسة بالذكاء الاصطناعي (Studia)", "AI Study Assistant (Studia)"),
    12: ("تقنية المعلومات والذكاء الاصطناعي", "Information Technology & Artificial Intelligence"),
    13: ("", ""),
    14: ("مرحلة التطوير والتكامل والاختبار", "Development, Integration & Testing"),
}
for row_idx, (ar, en) in metadata.items():
    # Use the center-left cell; merged-cell aliases are handled by python-docx.
    target = table.rows[row_idx].cells[1]
    replace_cell(target, ar, en, 8.5)

replace_cell(
    response_cell(table.rows[16]),
    "تم تطوير منصة ويب متكاملة لمساعدة الطلاب على تنظيم موادهم الدراسية والتفاعل معها بذكاء. شمل العمل ربط واجهة Next.js بخلفية FastAPI وقاعدة PostgreSQL، وتنفيذ إدارة الموضوعات والملاحظات، والمعلّم الذكي المعتمد على الاسترجاع المعزّز بالتوليد (RAG)، والاختبارات، والبطاقات التعليمية بالتكرار المتباعد، والاختبارات الشاملة، والخرائط الذهنية، والمدرب الدراسي، وسجل النشاط والتحليلات. كما تمت إضافة المصادقة الآمنة، وترحيلات قاعدة البيانات، والاختبارات الآلية، وتجهيزات النشر باستخدام Docker.",
    "A full-stack web platform was developed to help students organize and interact intelligently with their learning materials. The work connected a Next.js frontend to a FastAPI backend and PostgreSQL database, and implemented topics and notes management, a retrieval-augmented generation (RAG) AI tutor, quizzes, spaced-repetition flashcards, exams, mind maps, a study coach, activity history, and analytics. Secure authentication, database migrations, automated tests, and Docker-based deployment foundations were also added.",
    8.5,
)

replace_cell(
    response_cell(table.rows[19]),
    "اكتسبت خبرة عملية في بناء تطبيقات Full-Stack بهيكلية وحدات واضحة، وتصميم واجهات API باستخدام FastAPI، وإدارة البيانات والترحيلات عبر PostgreSQL وAlembic، وتطوير واجهات حديثة بـ Next.js وTypeScript. كما تعمقت المعرفة في RAG والبحث الهجين بالمتجهات والكلمات المفتاحية، وربط مزودي نماذج الذكاء الاصطناعي، وكتابة الاختبارات، وأساسيات الأمان والنشر.",
    "I gained practical experience in building modular full-stack applications, designing APIs with FastAPI, managing PostgreSQL data and Alembic migrations, and developing modern interfaces with Next.js and TypeScript. I also strengthened my knowledge of RAG, hybrid vector and keyword retrieval, AI-provider integration, automated testing, security, and deployment fundamentals.",
    8.5,
)

replace_cell(
    response_cell(table.rows[21]),
    "قدمت قيمة مضافة من خلال تحويل الفكرة إلى نظام قابل للاستخدام يربط أدوات الدراسة الأساسية في مكان واحد، مع جعل إجابات الذكاء الاصطناعي مرتبطة بملاحظات الطالب وملفاته ومصحوبة بالمصادر. كذلك ساهمت الهيكلية القابلة للتوسع والاختبارات الآلية في رفع موثوقية المشروع وتسهيل تطويره مستقبلًا.",
    "I added value by turning the concept into a usable system that brings essential study tools together in one place, while grounding AI answers in the student's own notes and documents with source references. The scalable architecture and automated tests also improved reliability and made future development easier.",
    8.5,
)

replace_cell(
    response_cell(table.rows[23]),
    "أبرز نتيجة هي الوصول إلى نسخة Full-Stack واسعة الوظائف تجمع تنظيم المحتوى، والمعلّم الذكي المدعوم بالمصادر، والاختبارات والبطاقات التعليمية، وتتبع التقدم في تجربة موحدة. أصبح المشروع جاهزًا للانتقال إلى مرحلة تحسين النشر وإجراء تجربة مستخدم تجريبية.",
    "The key outcome was reaching a feature-rich full-stack version that combines content organization, a source-grounded AI tutor, quizzes and flashcards, and progress tracking in one experience. The project is now ready for deployment hardening and supervised beta user testing.",
    8.5,
)

replace_cell(
    response_cell(table.rows[26]),
    "تمثلت أبرز التحديات في دمج عدد كبير من الوحدات مع الحفاظ على اتساق البيانات والصلاحيات، وضمان استرجاع معلومات دقيقة من ملاحظات المستخدم وملفاته، والتعامل مع اختلاف مزودي الذكاء الاصطناعي، إضافة إلى متطلبات الأمان والاختبار وتجهيز بيئة النشر.",
    "The main challenges were integrating many modules while maintaining consistent data and access control, retrieving accurate information from user notes and files, supporting different AI providers, and meeting security, testing, and deployment requirements.",
    8.5,
)

replace_cell(
    response_cell(table.rows[28]),
    "تم التعامل مع التحديات عبر تقسيم الخلفية إلى وحدات مستقلة تحتوي على المسارات والخدمات والمستودعات والنماذج، واستخدام بحث هجين يجمع pgvector وBM25 مع مسار بديل عند تعطل التضمينات. كما تم تطبيق جلسات آمنة وحدود للطلبات والتحقق من الملكية، وكتابة اختبارات لكل ميزة رئيسية، واستخدام ترحيلات قاعدة البيانات وDocker لضمان قابلية التشغيل والتوسع.",
    "The challenges were addressed by dividing the backend into independent route, service, repository, and model modules; using hybrid pgvector and BM25 retrieval with a fallback when embeddings fail; applying secure sessions, rate limits, and ownership checks; writing tests for each major feature; and using database migrations and Docker for repeatable operation and future scaling.",
    8.5,
)

replace_cell(response_cell(table.rows[31]), "5 / 5", "5 / 5", 10)

# Make all populated response rows grow naturally and avoid clipping.
for idx in [6, 8, 11, 12, 14, 16, 19, 21, 23, 26, 28, 31]:
    tr_pr = table.rows[idx]._tr.get_or_add_trPr()
    for node in list(tr_pr):
        if node.tag == qn("w:trHeight"):
            tr_pr.remove(node)

doc.save(OUTPUT)

# Basic package and content verification.
with ZipFile(OUTPUT) as zf:
    assert "word/document.xml" in zf.namelist()
check = Document(OUTPUT)
assert len(check.tables) == 1 and len(check.tables[0].rows) == 34
assert "مساعد الدراسة" in check.tables[0].rows[11].cells[1].text
assert "AI Study Assistant" in check.tables[0].rows[11].cells[1].text
print(OUTPUT)
